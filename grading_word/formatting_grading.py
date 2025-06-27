import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from checking_folder.checking_files import check_files_in_folder

# from docx.shared import WD_UNDERLINE.SINGLE, WD_UNDERLINE.DOUBLE, WD_UNDERLINE.WAVY
from grading_word.accuracy_grading import find_docs


def text_alignment(path):
    docx_list = find_docs(path)
    text_alignment_score = {}
    for docx in docx_list:
        filename = os.path.basename(docx)
        dirname = os.path.basename(os.path.dirname(docx))
        doc = Document(docx)
        # print(f"Body text in {filename}:")
        score = 0

        right_aligned_texts = ""  # Store right-aligned text as string

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                if paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    right_aligned_texts += paragraph.text + "\n"
                    # print(f"Right-aligned text: {paragraph.text}")

        # You can now use right_aligned_texts string for further processing
        if "yogyakarta, 07 juni 2018" in right_aligned_texts.lower():
            print(f"Found right-aligned text containing the date in {filename}")
            score += 10
        else:
            print(f"Date not found in right-aligned text in {filename}")

        text_alignment_score[dirname] = score

    return text_alignment_score


def bold_text(path):
    """Check for bold text in the document"""
    docx_list = find_docs(path)
    # text_alignment = {}
    bold_text_score = {}

    for docx in docx_list:
        # filename = os.path.basename(docx)
        dirname = os.path.basename(os.path.dirname(docx))

        doc = Document(docx)
        # print(f"Checking bold text in {filename}:")
        score = 0

        bold_texts = ""  # Store right-aligned text as string

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                for run in paragraph.runs:
                    if run.bold:
                        bold_texts += paragraph.text + "\n"
                        # print(f"Bold text found: {paragraph.text}")
                        break

        if "assalamualaikum wr. wb." in bold_texts.lower():
            # print(f"Found bold text containing the greeting in {filename}")
            score += 5
        if "wassalamualaikum wr. wb." in bold_texts.lower():
            # print(f"Found bold text containing the closing greeting in {filename}")
            score += 5

        bold_text_score[dirname] = score
        # print(f"Total bold text score for {filename}: {score}")

    return bold_text_score
    # return print(bold_texts)


def signature_alignment(path):
    docx_list = find_docs(path)
    signature_score_alignment = {}
    for docx in docx_list:
        dirname = os.path.basename(os.path.dirname(docx))

        filename = os.path.basename(docx)
        doc = Document(docx)

        score = 0

        underlined_texts = ""  # Collect underlined text

        # For paragraph text:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.underline:
                    underlined_texts += paragraph.text + "\n"
                    # print(f"Bold text found: {paragraph.text}")
                    break

        # For table cell text:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.underline:
                                underlined_texts += paragraph.text + "\n"
                                # print(f"Bold text found: {paragraph.text}")
                                break

        underlined_texts_lower = underlined_texts.lower()

        # Dr. Shofwatul Uyun, S.T.,M.Kom.

        if "dr. shofwatul uyun, s.t.,m.kom" in underlined_texts_lower:
            print(f"Found underlined text containing the signature in {filename}")
            score += 5

        signature_score_alignment[dirname] = score

    return signature_score_alignment


def find_pdf(path):
    # Get all files from the directory
    files = check_files_in_folder(path)

    pdf_scores = {}

    if not files:
        return pdf_scores

    pdf_files = [
        f
        for f in files
        if f.endswith(".pdf")
        and "word" in f.lower()
        and "cetak" in f.lower()
        and "-" in f
    ]

    for pdf in pdf_files:
        dirname = os.path.basename(
            os.path.dirname(pdf)
        )  # Get folder name from each file
        filename = os.path.basename(pdf)
        print(f"Found PDF file: {filename}")

        score = 0
        if "cetak" in filename.lower():
            score += 10
            print(f"Found PDF file with 'cetak' in the name: {filename}")
        else:
            print(f"No 'cetak' found in the PDF file name: {filename}")

        pdf_scores[dirname] = score

    return pdf_scores

    # print(f"Found {len(docx_files)} Word documents in the directory.")
    # return docx_files
    # else:
    #     print("Invalid path or not a Word document.")


# Grading function for formatting
def calculate_total_scores_formatting(path):
    """Calculate total scores for all documents"""
    body_alignment = text_alignment(path)
    bold_scores = bold_text(path)
    signature_score = signature_alignment(path)
    pdf_scores = find_pdf(path)

    total_scores = {}

    # Get all unique filenames
    all_files = (
        set(body_alignment.keys())
        | set(bold_scores.keys())
        | set(signature_score.keys())
        | set(pdf_scores.keys())
    )

    for dirname in all_files:
        total = (
            body_alignment.get(dirname, 0)
            + bold_scores.get(dirname, 0)
            + signature_score.get(dirname, 0)
            + pdf_scores.get(dirname, 0)
        )
        total_scores[dirname] = total
        # print(f"📄 {filename} — Total Score: {total}")

    return total_scores
