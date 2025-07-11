import os

from docx import Document

from grading_word.finding_docx import find_docs


def checking_image(path):
    docx_list = find_docs(path)
    has_image = []
    image_scores = {}  # Initialize image score
    for docx in docx_list:
        dirname = os.path.basename(os.path.dirname(docx))

        filename = os.path.basename(docx)
        doc = Document(docx)
        header = doc.sections[0].header

        found_image = False
        score = 0  # Initialize score for this document

        if len(header.tables) > 0:
            # print(f"Header table found in {filename}")
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if (
                            "graphic" in cell._element.xml
                            and "pic" in cell._element.xml
                        ):
                            print(f"Image found in header table of {filename}")
                            has_image.append(filename)
                            score += 5
                            found_image = True
                            break
                    if found_image:
                        break
                if found_image:
                    break

        # Always check paragraphs, not just in else clause
        if not found_image:
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    if "graphic" in run._element.xml and "pic" in run._element.xml:
                        print(f"Image found in header of {filename}")
                        has_image.append(filename)
                        found_image = True
                        score += 5
                        break
                if found_image:
                    break

        image_scores[dirname] = score

    # print(has_image)  # Print the list for debugging

    return image_scores  # Return dictionary for calculate_total_scores


# Checking if there is header inside the docx
# alse making sure that the header contains the required text
def header_info(path):
    docx_list = find_docs(path)
    header_score = {}
    for docx in docx_list:
        dirname = os.path.basename(os.path.dirname(docx))

        filename = os.path.basename(docx)
        doc = Document(docx)
        header = doc.sections[0].header
        score = 0

        # Check if document has multiple sections
        if len(doc.sections) > 1:
            second_header = doc.sections[1].header

            second_header_text = ""
            for p in second_header.paragraphs:
                second_header_text += p.text.strip()

            for table in second_header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        second_header_text += cell.text.strip() + " "

            if second_header_text.strip():
                # print(f"Second header found in {filename}: {second_header_text}")
                score -= 5

        # Get all text from header paragraphs
        header_text = ""
        for paragraph in header.paragraphs:
            header_text += paragraph.text

        # get all text from header tables
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    header_text += cell.text + " "

        if header_text.strip():
            # print(f"Header found in {filename}: {header_text.lower()}")
            if (
                "universitas islam negeri sunan kalijaga" in header_text.lower()
                and "pusat teknologi informasi dan pangkalan data"
                in header_text.lower()
            ):
                # print(f"Header contains required text: {header_text}")
                # print(f"✅ Header contains required text in {filename}")
                score += 10
            else:
                print(f"Header does not contain required text: {filename}")
        else:
            print(f"Header is empty in {filename}")

        header_score[dirname] = score

    return header_score


# Checking the body of the document and scoring it based on the required text
def body_info(path):
    docx_list = find_docs(path)
    body_score = {}

    # List of (phrase, tag)
    required_phrases = [
        ("dalam rangka upaya peningkatan pengetahuan dan keterampilan"),
        (
            "kami harapkan bapakibu dekan selaku pimpinan fakultas menugaskan 1 (satu) staff"
        ),
        (
            "demikian surat pemberitahuan ini kami sampaikan, atas perhatiannya kami ucapkan terima kasih."
        ),
    ]

    for docx in docx_list:
        filename = os.path.basename(docx)
        dirname = os.path.basename(os.path.dirname(docx))

        doc = Document(docx)

        body_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        lower_body = body_text.lower()

        score = 0

        for phrase in required_phrases:
            if phrase in lower_body:
                print(f"✅ Found required text in {filename}")
                score += 5

        body_score[dirname] = score

    return body_score


def signature_info(path):
    docx_list = find_docs(path)
    signature_score = {}
    for docx in docx_list:
        # filename = os.path.basename(docx)
        dirname = os.path.basename(os.path.dirname(docx))

        doc = Document(docx)

        score = 0

        # Get all text from header paragraphs
        signature_text = ""
        for paragraph in doc.paragraphs:
            signature_text += paragraph.text

        # get all text from header tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    signature_text += cell.text + " "

        signature_lower = signature_text.lower()

        # print(f"Signature text in {filename}: {signature_text.strip()}")
        if "hormat kami" in signature_lower or "wassalamualaikum" in signature_lower:
            score += 1
        if "kepala ptipd" in signature_lower:
            score += 1
        if "shofwatul uyun" in signature_lower:  # or regex for proper name pattern
            score += 1
        if (
            any(char.isdigit() for char in signature_lower)
            and "2006" in signature_lower
        ):
            score += 1
        if "hormat kami" in signature_lower and "kepala ptipd" in signature_lower:
            score += 1  # assuming structure is formal enough

        signature_score[dirname] = score

    return signature_score


def table_info(path):
    docx_list = find_docs(path)
    table_score = {}
    for docx in docx_list:
        filename = os.path.basename(docx)
        dirname = os.path.basename(os.path.dirname(docx))

        doc = Document(docx)

        score = 0

        table_text = []
        # table_text = ""
        # Check if there are tables in the Document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # table_text += cell.text + " "
                    table_text.append(cell.text.strip())

        # lowercase the table text for easier matching
        table_text_lower = [text.lower() for text in table_text]

        # if needed to debug the table text
        print(f"Table text in {filename}: {table_text_lower}")

        if "pembukaan dan sambutan-sambutan" in table_text_lower:
            print(f"✅ Found required text in table of {filename}")
            score += 5
        if "istirahat dan makan siang" in table_text_lower:
            print(f"✅ Found required text in table of {filename}")
            score += 5
        if "materi pengantar tik" in table_text_lower:
            print(f"✅ Found required text in table of {filename}")
            score += 5
        if "internet browsing dan email" in table_text_lower:
            print(f"✅ Found required text in table of {filename}")
            score += 5
        if "mengelola konten web" in table_text_lower:
            print(f"✅ Found required text in table of {filename}")
            score += 5
        if "membuat desain banner untuk web" in table_text_lower:
            print(f"✅ Found required text in table of {filename}")
            score += 5

        table_score[dirname] = score

    return table_score


# Grading function to calculate accuracy  score for all documents
def calculate_total_scores(path):
    """Calculate total scores for all documents"""
    image_score = checking_image(path)
    header_scores = header_info(path)
    body_scores = body_info(path)
    signature_scores = signature_info(path)
    table_scores = table_info(path)

    total_scores = {}

    # Get all unique filenames
    all_files = (
        set(image_score.keys())
        | set(header_scores.keys())
        | set(body_scores.keys())
        | set(signature_scores.keys())
        | set(table_scores.keys())
    )

    for dirname in all_files:
        total = (
            image_score.get(dirname, 0)
            + header_scores.get(dirname, 0)
            + body_scores.get(dirname, 0)
            + signature_scores.get(dirname, 0)
            + table_scores.get(dirname, 0)
        )
        total_scores[dirname] = total
        # print(f"📄 {filename} — Total Score: {total}")

    return total_scores
